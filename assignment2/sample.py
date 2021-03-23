import sys
import docx
import re
import os
doc = docx.Document('C://Users//91720//Documents//ds_reg_exp_proj//assignment2//ulysses.docx')
filename = doc

def count_docx(file_name):
    document = docx.Document(file_name)

    newparatextlist = []
    for paratext in document.paragraphs:
        newparatextlist.append(paratext.text)
    
    return len(re.findall(r'\w+', '\n'.join(newparatextlist)))

if __name__ == '__main__':
    extensions = {
        'docx' : count_docx,
        'doc' : count_docx,
    }
    # file_name, file_extension = os.path.splitext(sys.argv[1])
    # print(extensions[file_extension.lower().replace('.','')](sys.argv[1]))
    
print(count_docx(filename))